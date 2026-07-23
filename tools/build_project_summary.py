#!/usr/bin/env python3
"""Assemble the project-summary Word document (docs/feature_summary_table.docx).

The document has three parts, in order:
  1. A short narrative review (<= ~1/2 page) of what the project did and the current state
     of the annotations, written for expert scientists who are not necessarily AI experts.
  2. Table 1 - the class x level feature-summary table (reused from build_feature_summary).
  3. One to three figures from plotFactorScores (matlab/figures/), each with a caption.

Headline numbers are pulled live from schema/channel_template.json, analysis/corpus_stats.json,
and analysis/factor_reducibility.json (written by the MATLAB factor analysis) so the prose stays
accurate as the corpus grows. See docs/summarize_project.md for the reusable instruction this
implements ("summarize the project now").

    python3 tools/build_project_summary.py       # -> docs/feature_summary_table.docx

Requires python-docx.
"""
from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT / "tools"))
import build_feature_summary as bfs   # reuse STRUCTURE, load_counts, cell_text, shade

OUT = ROOT / "docs" / "feature_summary_table.docx"
FIGDIR = ROOT / "matlab" / "figures"

# Figures to embed (file stem, caption). Chosen to show the annotations most clearly.
FIGURES = [
    ("01_factor_timeseries",
     "Figure 1. Second-by-second factor time courses for one audiovisual clip (Kung Fury). Each "
     "horizontal trace is one of the {n_factors} per-model factors, colored by feature domain "
     "(reds = emotion/affect, cyans = auditory, blues/purples = visual, orange = language, "
     "pink = social, green = situational). The annotations vary continuously and coherently as "
     "the scene unfolds, which is the raw material for relating stimulus content to brain "
     "responses."),
    ("04_correlation_matrix",
     "Figure 2. Correlation matrix of all {n_factors} factors across the entire corpus "
     "({n_seconds:,} s), ordered and color-labeled by category. Near-zero correlations inside each "
     "outlined diagonal block show that the factors within a model are largely non-redundant; "
     "off-diagonal structure reveals genuine cross-modal coupling (e.g., among the visual "
     "embedding, action, and depth factors)."),
    ("05_tsne",
     "Figure 3. Two-dimensional t-SNE layout of the {n_factors} factors, each point one factor "
     "colored by category. Factors from the same model or sensory domain tend to group together, "
     "while emotion and interpretable factors spread across the space, reflecting variance they "
     "share with several modalities."),
    ("08_umap_annotations",
     "Figure 4. UMAP embedding of all {n_vars:,} individual feature-variable annotations (not the "
     "factors), each point one variable colored by feature domain (legend). The variables separate "
     "cleanly into three modality masses — visual, auditory, and language (now the largest, "
     "dominated by the Qwen3 and Llama-3.1 language embeddings) — with the smaller social, "
     "situational, and affective feature sets between them, showing that the annotation space is "
     "organized primarily by modality."),
]


def facts() -> dict:
    tmpl = json.loads((ROOT / "schema" / "channel_template.json").read_text())
    n_channels = tmpl.get("n_channels", len(tmpl["channels"]))
    n_vars = 0
    for c in tmpl["channels"]:
        dt = c.get("dtype")
        if dt == "vector":
            n_vars += int(c.get("dim") or 0)
        elif dt in ("scalar", "bool", "event"):
            n_vars += 1
    f = {"n_channels": n_channels, "n_variables": n_vars,
         "n_stimuli": 83, "audiovisual": 53, "audio_only": 29, "text_only": 1,
         "total_minutes": 470.6}
    cs = ROOT / "analysis" / "corpus_stats.json"
    if cs.exists():
        s = json.loads(cs.read_text())
        f["n_stimuli"] = s.get("n_stimuli", f["n_stimuli"])
        f["total_minutes"] = s.get("total_minutes", f["total_minutes"])
        for k in ("audiovisual", "audio_only", "text_only"):
            f[k] = s.get("by_modality", {}).get(k, f[k])
    red = ROOT / "analysis" / "factor_reducibility.json"
    f["reducibility"] = json.loads(red.read_text()) if red.exists() else None
    return f


def para(doc, text, size=11, bold=False, italic=False, color=None, after=8, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.bold, r.italic = bold, italic
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor(*color)
    p.paragraph_format.space_after = Pt(after)
    return p


def add_table1(doc, f):
    nvar = bfs.load_counts()
    para(doc, "Table 1. Feature groups by class (rows) and level of abstraction (columns).",
         size=10.5, bold=True, after=2)
    para(doc, "Numbers in parentheses are the count of variables in each group (a scalar/flag = 1; "
              "a vector = its dimensionality). Low-level = raw physical/perceptual signal; "
              "mid-level = perceptual organization; high-level = semantic/conceptual.",
         size=8.5, italic=True, color=(0x55, 0x55, 0x55), after=6)

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for cell, text in zip(hdr, ["Feature class"] + bfs.LEVELS):
        pr = cell.paragraphs[0]; pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = pr.add_run(text); rr.bold = True; rr.font.size = Pt(8.5)
        rr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        bfs.shade(cell, "334155")
    for cls, levels in bfs.STRUCTURE.items():
        cells = table.add_row().cells
        cp = cells[0].paragraphs[0]
        cr = cp.add_run(cls); cr.bold = True; cr.font.size = Pt(8.5)
        cr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        bfs.shade(cells[0], bfs.CLASS_HUE[cls])
        for i, level in enumerate(bfs.LEVELS, start=1):
            rr = cells[i].paragraphs[0].add_run(bfs.cell_text(levels.get(level, []), nvar))
            rr.font.size = Pt(8.5)
    widths = [Inches(0.9), Inches(1.83), Inches(1.83), Inches(1.84)]
    for row in table.rows:
        for cell, w in zip(row.cells, widths):
            cell.width = w


def reducibility_phrases(f):
    """Build 'irreducible' and 'reducible' example clauses from the live JSON."""
    red = f.get("reducibility")
    if not red:
        return ("several high-dimensional model outputs", "a few low-dimensional feature sets")
    cats = sorted(red["categories"], key=lambda c: c["pc1_pct"])
    pretty = {"action_posteriors": "the 400-way action posteriors",
              "dino_embedding": "the DINOv2 image embedding",
              "siglip_embedding": "the SigLIP2 image embedding",
              "audioset_tags": "the 527-way AudioSet sound tags",
              "emonet": "the EmoNet visual-emotion schemas",
              "clap_embedding": "the CLAP audio embedding",
              "chroma": "musical pitch-class (chroma)",
              "text_sentiment": "text sentiment",
              "mfcc": "MFCC timbre",
              "qwen3": "the Qwen3 language embedding",
              "llama_ar": "the Llama-3.1 language embedding"}
    low = [c for c in cats if c["pc1_pct"] <= 7][:5]
    high = [c for c in reversed(cats) if c["pc1_pct"] >= 50][:2]
    low_txt = "; ".join(f"{pretty.get(c['category'], c['category'])} ({c['pc1_pct']:.0f}%)" for c in low)
    high_txt = "; ".join(f"{pretty.get(c['category'], c['category'])} ({c['pc1_pct']:.0f}%)" for c in high)
    return low_txt, high_txt


def main():
    f = facts()
    n_seconds = int(round(f["total_minutes"] * 60))
    hours = f["total_minutes"] / 60
    n_factors = 176
    if f.get("reducibility"):
        n_seconds = f["reducibility"].get("n_timepoints", n_seconds)
    low_txt, high_txt = reducibility_phrases(f)

    doc = Document()
    sec = doc.sections[0]
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, m, Inches(0.9))
    doc.styles["Normal"].font.name = "Calibri"

    para(doc, "Narrative Feature Extraction — Project Summary", size=16, bold=True, after=6)

    para(doc,
         "Naturalistic neuroimaging increasingly uses rich, dynamic stimuli — movies and spoken "
         "stories — that engage perception, language, emotion, and social cognition as they unfold "
         "in time. Making sense of brain responses to such stimuli requires quantitative, "
         "second-by-second descriptions of what is on the screen and in the soundtrack. We built "
         "infrastructure to produce these descriptions automatically and reproducibly. We first "
         "performed a broad scoping review of computational models for visual, auditory, and "
         "language annotation, surveying roughly 146 tools across 15 feature subclasses — from "
         "interpretable features built with classical machine learning and signal processing to "
         "large ‘foundation’ and multimodal transformer models. We preferred models with "
         "large, well-documented training sets, open weights (so annotations can be inspected, "
         "shared, and reproduced), and best-in-class validated performance. From these we selected "
         "24 feature extractors — about half interpretable classical-ML / signal-processing "
         "features (e.g., low-level image statistics, acoustic descriptors, psycholinguistic word "
         "norms, GPT-2 word surprisal, optical flow, event segmentation) and about half large "
         "open-weight models with expressive learned feature spaces (e.g., the DINOv2 and SigLIP2 "
         "vision transformers, CLAP audio–language embeddings, VideoMAE action recognition, "
         "EmoNet visual-emotion schemas, a Qwen2.5 vision–language model, and — for a rich "
         "representation of what is said — Qwen3 sentence embeddings and Llama-3.1-8B autoregressive "
         "hidden states of the dialogue and narration). The features are "
         f"organized into 6 broad domains — visual, auditory, language, social, emotional (affect), "
         f"and situational/schema-level — each spanning low-, mid-, and high-level descriptions "
         f"(Table 1).")

    para(doc,
         f"Applying these models, we extracted about {f['n_variables']:,} feature-variable time "
         f"series (from {f['n_channels']} annotation channels) across {f['n_stimuli']} movie clips "
         f"and stories ({f['audiovisual']} audiovisual, {f['audio_only']} spoken-audio, "
         f"{f['text_only']} text), yielding roughly {n_seconds:,} seconds (~{hours:.1f} h) of "
         "annotations sampled once per second on a common time grid. Features that do not apply to "
         "a stimulus (e.g., visual features for an audio-only story) are stored as explicit nulls, "
         "so the whole corpus stacks into a single rectangular matrix; embedding the full set of "
         "variables shows that this annotation space is organized primarily by modality, with the "
         "dense LLM language embeddings forming its largest, cleanly separated cluster (Figure 4). "
         "An interactive browser lets users visualize and manually check each "
         "feature’s distribution and quality. To "
         "characterize the structure of the annotation space, we performed exploratory factor "
         "analysis within each domain/model. Some feature sets proved essentially irreducible — the "
         f"first principal factor of {low_txt} each explained only a small share of the variance, "
         "indicating genuinely high-dimensional, distributed representations that cannot be "
         "collapsed without losing information. Others were highly compressible — a single component "
         f"captured most of the variance for {high_txt}. Reducing each model to its leading factors "
         f"yields {n_factors} compact, labeled regressors (Figures 1–3) suited to encoding "
         "models that relate stimulus content to brain activity.")

    # Table 1 in its own landscape section (wide), then back to portrait for figures.
    lsec = doc.add_section(WD_SECTION.NEW_PAGE)
    lsec.orientation = WD_ORIENT.LANDSCAPE
    lsec.page_width, lsec.page_height = Inches(11), Inches(8.5)
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(lsec, m, Inches(0.7))
    add_table1(doc, f)

    psec = doc.add_section(WD_SECTION.NEW_PAGE)
    psec.orientation = WD_ORIENT.PORTRAIT
    psec.page_width, psec.page_height = Inches(8.5), Inches(11)
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(psec, m, Inches(0.9))
    para(doc, "Figures", size=13, bold=True, after=6)
    ctx = {"n_factors": n_factors, "n_seconds": n_seconds, "n_vars": f["n_variables"]}
    for stem, caption in FIGURES:
        png = FIGDIR / f"{stem}.png"
        if not png.exists():
            para(doc, f"[missing figure: {png.name} — run plotFactorScores in MATLAB]",
                 size=9, italic=True, color=(0x99, 0x33, 0x33))
            continue
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(png), width=Inches(6.5))
        para(doc, caption.format(**ctx), size=9, italic=True,
             color=(0x33, 0x33, 0x33), after=14)

    doc.save(OUT)
    print(f"wrote {OUT}")
    print(f"  {f['n_variables']:,} variables, {f['n_stimuli']} stimuli, {n_seconds:,} s, "
          f"{n_factors} factors; {sum(1 for s,_ in FIGURES if (FIGDIR/f'{s}.png').exists())} figures")


if __name__ == "__main__":
    main()
