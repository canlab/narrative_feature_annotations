#!/usr/bin/env python3
"""Generate the one-page feature-summary table as a Word .docx (for import into Google Docs).

Rows = the six feature classes; columns = level of abstraction (low / mid / high). Each cell
lists the feature groups at that level with a variable count in parentheses (a scalar/flag
counts as 1, a vector counts as its dimensionality), computed from schema/channel_template.json
so the docx stays in sync with the pipeline. Mirrors the "Feature summary table" in
docs/FEATURE_MAP.md. Landscape + narrow margins so it fits on a single page.

    python3 tools/build_feature_summary.py     # -> docs/feature_summary_table.docx

Requires python-docx (`pip install python-docx`).
"""
from __future__ import annotations

import json
from collections import OrderedDict
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "docs" / "feature_summary_table.docx"

# Level order and a header color per class (matches the feature-map hues).
LEVELS = ["Low-level", "Mid-level", "High-level"]
CLASS_HUE = {"Visual": "6366F1", "Audio": "06B6D4", "Language": "F59E0B",
             "Social": "EC4899", "Situation": "10B981", "Affect": "EF4444"}

# class -> level -> list of (label, [channel paths]). Counts are filled from the template.
STRUCTURE = OrderedDict([
    ("Visual", {
        "Low-level": [
            ("image statistics", ["visual/low_level_static/" + s for s in
                ["luminance", "rms_contrast", "r_mean", "g_mean", "b_mean", "hue_mean",
                 "sat_mean", "val_mean", "colorfulness", "edge_density", "entropy", "fft_slope"]]),
            ("optical-flow motion", ["visual/dynamic_motion/flow_magnitude",
                "visual/dynamic_motion/camera_motion", "visual/dynamic_motion/residual_motion"]),
            ("shot cuts", ["visual/dynamic_motion/cut", "visual/dynamic_motion/shot_index"]),
        ],
        "Mid-level": [
            ("saliency & depth", ["visual/saliency_aesthetics_depth/" + s for s in
                ["saliency_mean", "saliency_peak", "saliency_entropy", "salient_area_frac",
                 "depth_mean", "depth_range", "depth_entropy", "foreground_frac"]]),
            ("faces & bodies", ["visual/faces_bodies_gaze/" + s for s in
                ["faces_present", "n_faces", "face_det_prob", "max_face_frac",
                 "pose_present", "n_persons", "mean_kp_score"]]),
        ],
        "High-level": [
            ("SigLIP semantics — embedding + scene probes",
                ["visual/high_level_static/siglip_embedding", "visual/high_level_static/siglip_probe"]),
            ("DINOv2 embedding", ["visual/high_level_static/dino_embedding"]),
            ("action recognition", ["visual/action/action_posteriors"]),
        ],
    }),
    ("Audio", {
        "Low-level": [
            ("acoustics — loudness / pitch / timbre / MFCC / chroma",
                ["audio/low_level/" + s for s in
                 ["rms", "f0", "tempo", "onset_strength", "zcr", "spectral_centroid",
                  "spectral_bandwidth", "spectral_rolloff", "spectral_flatness", "mfcc", "chroma"]]),
        ],
        "Mid-level": [
            ("speech presence & rate", ["audio/speech/speech_present", "audio/speech/word_rate"]),
        ],
        "High-level": [
            ("AudioSet sound tags", ["audio/high_level/audioset_tags"]),
            ("CLAP semantics — embedding + probes",
                ["audio/high_level/clap_embedding", "audio/high_level/clap_probe"]),
            ("vocal affect V/A/D", ["audio/speech/voice_valence",
                "audio/speech/voice_arousal", "audio/speech/voice_dominance"]),
        ],
    }),
    ("Language", {
        "Low-level": [
            ("word frequency & length", ["language/lexical/freq_zipf", "language/lexical/word_length"]),
        ],
        "Mid-level": [
            ("lexical norms — valence / arousal / dominance, concreteness, AoA",
                ["language/lexical/" + s for s in
                 ["valence", "arousal", "dominance", "concreteness", "aoa"]]),
            ("syntactic complexity", ["language/syntax/" + s for s in
                ["content_frac", "noun_frac", "verb_frac", "tree_depth", "mean_dep_distance"]]),
        ],
        "High-level": [
            ("word surprisal & next-word entropy",
                ["language/lexical/surprisal", "language/lexical/entropy"]),
        ],
    }),
    ("Social", {
        "Low-level": [],
        "Mid-level": [
            ("agent count & closest-pair distance", ["social/n_agents", "social/min_pair_distance"]),
        ],
        "High-level": [
            ("interaction type & social dominance", ["social/interaction_type", "social/dominance"]),
        ],
    }),
    ("Situation", {
        "Low-level": [],
        "Mid-level": [],
        "High-level": [
            ("scene description / setting / indoor–outdoor",
                ["situation/scene_description", "situation/indoor_outdoor", "situation/setting"]),
            ("event structure — boundaries & event id",
                ["situation/event_boundary", "situation/event_id"]),
        ],
    }),
    ("Affect", {
        "Low-level": [],
        "Mid-level": [
            ("facial affect — 8 expressions + valence + arousal",
                ["affect/depicted/face_emotion", "affect/depicted/face_valence",
                 "affect/depicted/face_arousal"]),
        ],
        "High-level": [
            ("EmoNet image emotion", ["affect/depicted/emonet"]),
            ("text emotion & sentiment", ["affect/depicted/text_emotion",
                "affect/depicted/text_sentiment", "affect/depicted/text_sentiment_polarity"]),
            ("VLM depicted emotion", ["affect/depicted/vlm_emotion",
                "affect/depicted/vlm_valence", "affect/depicted/vlm_arousal"]),
        ],
    }),
])


def load_counts():
    t = json.loads((ROOT / "schema" / "channel_template.json").read_text())
    ch = {c["path"]: c for c in t["channels"]}
    def nvar(paths):
        return sum(ch[p]["dim"] if ch[p]["dtype"] == "vector" else 1 for p in paths)
    return nvar


def cell_text(groups, nvar):
    if not groups:
        return "—"
    return "; ".join(f"{label} ({nvar(paths)})" for label, paths in groups)


def shade(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.makeelement(qn("w:shd"), {qn("w:val"): "clear", qn("w:fill"): hex_fill})
    tcPr.append(shd)


def set_widths(table, widths):
    table.autofit = False
    table.allow_autofit = False
    for row in table.rows:
        for cell, w in zip(row.cells, widths):
            cell.width = w


def main():
    nvar = load_counts()
    doc = Document()

    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = Inches(11), Inches(8.5)
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, m, Inches(0.5))

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(9)

    h = doc.add_paragraph()
    r = h.add_run("Narrative Feature Annotation — Summary by class and level")
    r.bold = True
    r.font.size = Pt(14)
    sub = doc.add_paragraph()
    sr = sub.add_run(
        "Feature groups by class (rows) and level of abstraction (columns). Numbers in "
        "parentheses = count of variables (each scalar/flag = 1; each vector = its "
        "dimensionality). Low-level = raw physical/perceptual signal; mid-level = "
        "perceptual organization; high-level = semantic/conceptual.")
    sr.italic = True
    sr.font.size = Pt(8)
    sr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = table.rows[0].cells
    for cell, text in zip(hdr, ["Feature class"] + LEVELS):
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = p.add_run(text)
        rr.bold = True
        rr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade(cell, "334155")

    for cls, levels in STRUCTURE.items():
        cells = table.add_row().cells
        cp = cells[0].paragraphs[0]
        cr = cp.add_run(cls)
        cr.bold = True
        cr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade(cells[0], CLASS_HUE[cls])
        for i, level in enumerate(LEVELS, start=1):
            cells[i].paragraphs[0].add_run(cell_text(levels.get(level, []), nvar))

    set_widths(table, [Inches(1.1)] + [Inches(2.95)] * 3)

    total = sum(nvar(paths) for levels in STRUCTURE.values()
                for groups in levels.values() for _, paths in groups)
    foot = doc.add_paragraph()
    fr = foot.add_run(
        f"Totals: {sum(1 for _ in STRUCTURE)} classes, "
        f"{sum(len(g) for lv in STRUCTURE.values() for g in lv.values())} feature groups, "
        f"{total} variables (dominated by the SigLIP/DINOv2/CLAP embeddings and the "
        f"AudioSet/Kinetics taxonomies). Source: schema/channel_template.json.")
    fr.italic = True
    fr.font.size = Pt(7.5)
    fr.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    doc.save(OUT)
    print(f"wrote {OUT}  ({total} variables across "
          f"{sum(len(g) for lv in STRUCTURE.values() for g in lv.values())} groups)")


if __name__ == "__main__":
    main()
